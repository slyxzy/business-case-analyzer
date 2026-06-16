"""
comments.py
Extracts reviewer comments and body text from .docx files.
"""
import zipfile
from docx import Document
from lxml import etree

WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _load_comments_xml(filepath: str) -> bytes | None:
    """Read the comments XML part from the .docx ZIP package."""
    try:
        with zipfile.ZipFile(filepath, "r") as archive:
            if "word/comments.xml" in archive.namelist():
                return archive.read("word/comments.xml")

            rels_path = "word/_rels/document.xml.rels"
            if rels_path not in archive.namelist():
                return None

            rels = archive.read(rels_path).decode("utf-8")
            start = rels.find("Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments\"")
            if start < 0:
                return None
            target_start = rels.rfind("Target=\"", 0, start)
            if target_start < 0:
                return None
            target_start += len("Target=\"")
            target_end = rels.find("\"", target_start)
            if target_end < 0:
                return None
            target = rels[target_start:target_end]
            target_path = target if target.startswith("word/") else f"word/{target}" if not target.startswith("/") else target.lstrip("/")
            if target_path in archive.namelist():
                return archive.read(target_path)
    except Exception:
        return None
    return None


def extract_comments(filepath: str) -> list[dict]:
    """Return a list of {author, date, text} dicts from a .docx file."""
    comments = []
    root = None

    try:
        doc = Document(filepath)
        if hasattr(doc.part, "comments_part") and doc.part.comments_part is not None:
            root = doc.part.comments_part._element
    except Exception:
        root = None

    if root is None:
        xml_bytes = _load_comments_xml(filepath)
        if xml_bytes:
            root = etree.fromstring(xml_bytes)

    if root is None:
        print(f"Warning: no comments found or unsupported .docx comments in {filepath}")
        return comments

    for comment in root.iter(f"{WNS}comment"):
        author = comment.get(f"{WNS}author", "Unknown")
        date = comment.get(f"{WNS}date", "")[:10]
        text = "".join(
            t.text for t in comment.iter(f"{WNS}t") if t.text
        ).strip()
        if text:
            comments.append({"author": author, "date": date, "text": text})
    return comments


def extract_body_text(filepath: str) -> str:
    """Return the full body text of a .docx file."""
    doc = Document(filepath)
    text = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)

    return "\n".join(text)


def format_comments_for_llm(comments: list[dict]) -> str:
    """Format comments into a readable string for the LLM context."""
    if not comments:
        return "No reviewer comments found."
    lines = [f"[{c['date']}] {c['author']}: {c['text']}" for c in comments]
    return "\n".join(lines)